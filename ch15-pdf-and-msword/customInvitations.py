#! python3
# customInvitations.py - The program extracts names from a text file and for each name 
# makes a custom invitations for an event in MS Word. Each invitation takes up one page in MS Word.
#
# NOTES: Run this program in the same location as the other two required files. Those files are:
#         1. A text file called 'guests.txt' which contains names to be used in the invitation. One name per line.
#         2. A Word document with custom styles 'party', 'party_name' and 'party_date' created. 
#            The program can't create new styles and must load them from an exsiting Word document.
#         The result will be saved in the same location as the program, in a file called 'invitations.docx'.

import docx

doc = docx.Document('invite_template.docx')
doc._body.clear_content() # Clear all content from the Word doc. We just need the styles from it.

# Read the names from the txt file and save them into a list.
f = open('guests.txt')
names = f.readlines()
names = [name.rstrip() for name in names] # Strip the newline character from each name.
f.close()

# Add content and styles to the Word doc.
for name in names:
    doc.add_paragraph('It would be a pleasure to have the company of')
    doc.paragraphs[-1].style = 'party'
    doc.add_paragraph(name)
    doc.paragraphs[-1].style = 'party_name'
    doc.add_paragraph('at 11010 Memory Lane on the Evening of')
    doc.paragraphs[-1].style = 'party'
    doc.add_paragraph('April 1st')
    doc.paragraphs[-1].style = 'party_date'
    doc.add_paragraph('at 7 o\'clock')
    doc.paragraphs[-1].style = 'party'
    
    # Add a pagebreak at the end of each invite.
    if names.index(name) != len(names) - 1:
        doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)

doc.save('invites.docx')